import requests
import json
import os
from datetime import datetime
import tkinter as tk
from tkinter import messagebox, ttk, scrolledtext
from PIL import Image, ImageTk
import subprocess
from docx import Document
from docx.shared import Pt

# 读取配置文件
with open('./config/ai_config.json', 'r', encoding='utf-8') as file:
    config_data = json.load(file)
    ENDPOINT = config_data.get('Azure_setting').get("path").get("azure_chat_api")
    API_KEY = config_data.get('Azure_setting').get("key").get("azure_api_key")
    default_prompt = config_data.get('Azure_setting').get("prompt")
    interaction_prompts = config_data.get('Azure_setting').get("interaction_prompt")

with open('./config/paddlex_config.json', 'r', encoding='utf-8') as file:
    pdx_data = json.load(file)
    #ocr_result_dir = pdx_data.get('user_setting').get('output_path')
    ocr_result_dir = 'datas\ocr_result'

headers = {
    "Content-Type": "application/json",
    "api-key": API_KEY,
}

def resize_image_to_fill(image_path, canvas_width, canvas_height):
    image = Image.open(image_path)
    original_width, original_height = image.size

    # 计算适合画布的缩放比例，保持图片的宽高比
    scale = max(canvas_width / original_width, canvas_height / original_height)
    
    # 计算新的图片尺寸
    new_width = int(original_width * scale)
    new_height = int(original_height * scale)

    # 调整图片大小
    resized_image = image.resize((new_width, new_height), Image.LANCZOS)
    return resized_image

def save_results_to_docx(log_entry, selected_prompts, log_file_name):
    doc_file_path = os.path.join("datas", "result", log_file_name + ".docx")
    document = Document()
    document.add_heading('AI Interaction Results', level=1)
    document.add_paragraph(f"Timestamp: {log_entry['timestamp']}")
    if "translate_result" in log_entry:
        document.add_paragraph(f"Translation Result: {log_entry['translate_result']}")
    for prompt_key in selected_prompts:
        output_key = f"{prompt_key}_result"
        if output_key in log_entry:
            document.add_paragraph(f"{prompt_key} Result: {log_entry[output_key]}")
    
    # 设置字体大小
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)
    
    document.save(doc_file_path)

def process_files(selected_prompts, include_translation, output_text_widget, custom_prompt_text, save_to_file):
    for filename in os.listdir(ocr_result_dir):
        if filename.endswith('.json'):
            file_path = os.path.join(ocr_result_dir, filename)
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                rec_text_list = data.get('rec_text', [])
                text_input = ' '.join([f"[{text}]" for text in rec_text_list])

                log_entry = {"timestamp": datetime.now().isoformat(), "user_input": text_input}

                if include_translation:
                    prompt_input = default_prompt
                    payload = {
                        "messages": [
                            {
                                "role": "system",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": prompt_input
                                    }
                                ]
                            },
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": text_input
                                    }
                                ]
                            }
                        ],
                        "temperature": 0.2,
                        "top_p": 0.7,
                        "max_tokens": 800
                    }
                    
                    # 调试代码：打印请求负载
                    print(f"Request payload: {json.dumps(payload, ensure_ascii=False, indent=4)}")

                    try:
                        response = requests.post(ENDPOINT, headers=headers, json=payload)
                        response.raise_for_status()
                        
                        # 调试代码：打印响应状态码和内容
                        print(f"Response status code: {response.status_code}")
                        print(f"Response content: {response.content.decode('utf-8')}")
                        
                        result = response.json()
                        log_entry["translate_result"] = result['choices'][0]['message']['content']
                        output_text_widget.insert(tk.END, f"翻译结果: {log_entry['translate_result']}\n")
                    except requests.RequestException as e:
                        messagebox.showerror("错误", f"请求失败。错误: {e}")
                        return

                for prompt_key in selected_prompts:
                    prompt_input = interaction_prompts.get(prompt_key, "")
                    if prompt_key == "custom_prompt" and custom_prompt_text:
                        prompt_input = custom_prompt_text

                    payload = {
                        "messages": [
                            {
                                "role": "system",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": prompt_input
                                    }
                                ]
                            },
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": text_input
                                    }
                                ]
                            }
                        ],
                        "temperature": 0.2,
                        "top_p": 0.7,
                        "max_tokens": 800
                    }
                    
                    # 调试代码：打印请求负载
                    print(f"Request payload: {json.dumps(payload, ensure_ascii=False, indent=4)}")

                    try:
                        response = requests.post(ENDPOINT, headers=headers, json=payload)
                        response.raise_for_status()
                        
                        # 调试代码：打印响应状态码和内容
                        print(f"Response status code: {response.status_code}")
                        print(f"Response content: {response.content.decode('utf-8')}")
                        
                        result = response.json()
                        output_key = f"{prompt_key}_result"
                        log_entry[output_key] = result['choices'][0]['message']['content']
                        output_text_widget.insert(tk.END, f"{prompt_key} 结果: {log_entry[output_key]}\n")
                    except requests.RequestException as e:
                        messagebox.showerror("错误", f"请求失败。错误: {e}")
                        return

                log_file_name = os.path.splitext(filename)[0]
                log_file_path = os.path.join("datas", "result", log_file_name + ".json")
                os.makedirs(os.path.dirname(log_file_path), exist_ok=True)
                with open(log_file_path, "w", encoding="utf-8") as log_file:
                    json.dump(log_entry, log_file, ensure_ascii=False, indent=4)

                if save_to_file:
                    save_results_to_docx(log_entry, selected_prompts, log_file_name)

def create_gui():
    root = tk.Tk()
    root.title("无界视觉Ai交互-beta")
    root.geometry("600x550")

    # 创建画布用于放置背景图片
    canvas = tk.Canvas(root, width=600, height=550)
    canvas.pack(fill="both", expand=True)

    # 加载和设置背景图片
    background_image = resize_image_to_fill("src/assets/background/background2.jpg", 600, 550)
    background_photo = ImageTk.PhotoImage(background_image)
    canvas.create_image(0, 0, image=background_photo, anchor="nw")

    # 设置标题
    canvas.create_text(300, 30, text="无界视觉Ai交互-beta", font=("Arial", 16, "bold"), fill="white")

    # 创建选项和按钮
    translation_var = tk.BooleanVar()
    canvas.create_window(300, 80, window=tk.Checkbutton(root, text="翻译", variable=translation_var, bg="lightgray"))

    check_vars = {}
    y_offset = 120
    custom_prompt_entry = None
    for prompt_key in interaction_prompts.keys():
        var = tk.BooleanVar()
        check_vars[prompt_key] = var
        if prompt_key == "custom_prompt":
            canvas.create_window(300, y_offset, window=tk.Checkbutton(root, text="自定义提示", variable=var, bg="lightgray"))
            custom_prompt_entry = tk.Entry(root, width=40)
            canvas.create_window(300, y_offset + 30, window=custom_prompt_entry)
            y_offset += 60
        else:
            canvas.create_window(300, y_offset, window=tk.Checkbutton(root, text=prompt_key, variable=var, bg="lightgray"))
            y_offset += 30

    # 添加“输出文件”选项
    save_to_file_var = tk.BooleanVar()
    canvas.create_window(300, y_offset, window=tk.Checkbutton(root, text="输出文件", variable=save_to_file_var, bg="lightgray"))
    y_offset += 30

    # 添加“进行OCR识别”选项
    ocr_var = tk.BooleanVar(value=True)  # 默认勾选
    canvas.create_window(300, y_offset, window=tk.Checkbutton(root, text="进行OCR识别（除非你已经有过识别记录，否则请勾选该项选项）", variable=ocr_var, bg="lightgray"))
    y_offset += 30

    # 添加输出文本框
    output_text = scrolledtext.ScrolledText(root, width=70, height=10, wrap=tk.WORD)
    canvas.create_window(300, 400, window=output_text)

    def on_process():
        selected_prompts = [key for key, var in check_vars.items() if var.get()]
        if not selected_prompts and not translation_var.get():
            messagebox.showwarning("警告", "请选择至少一个处理选项。")
            return

        output_text.delete('1.0', tk.END)  # 清空输出框

        # 如果选择了进行OCR识别，先运行pocr.py
        if ocr_var.get():
            try:
                subprocess.run(['python', './pocr.py'], check=True)
            except subprocess.CalledProcessError as e:
                messagebox.showerror("错误", f"OCR识别失败。错误: {e}")
                return

        custom_prompt_text = custom_prompt_entry.get() if custom_prompt_entry else ""
        process_files(selected_prompts, translation_var.get(), output_text, custom_prompt_text, save_to_file_var.get())
        messagebox.showinfo("信息", "处理完成！")

    # 使用 ttk.Button 创建带有样式的按钮
    style = ttk.Style()
    style.configure("TButton", font=("Arial", 10, "bold"), relief="raised", padding=6)

    process_button = ttk.Button(root, text="处理", command=on_process, style="TButton")
    canvas.create_window(300, 500, window=process_button)

    root.mainloop()

create_gui()